from __future__ import annotations

from pathlib import Path
from typing import Dict


def read_docx_text(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    texts: list[str] = []
    for p in doc.paragraphs:
        texts.append(p.text)
    return "\n".join(texts)


def write_docx_with_replacements(input_path: Path, output_path: Path, replacements: Dict[str, str]) -> None:
    """Very simple paragraph-level replacements. Formatting across runs may be simplified.
    """
    from docx import Document  # type: ignore

    doc = Document(str(input_path))
    for para in doc.paragraphs:
        text = para.text
        new_text = _apply_replacements(text, replacements)
        if new_text != text:
            # replace the entire paragraph runs to preserve simple formatting
            for r in para.runs:
                r.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _apply_replacements(text: str, repl: Dict[str, str]) -> str:
    for k, v in repl.items():
        if k:
            text = text.replace(k, v)
    return text

